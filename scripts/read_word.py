#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档读取工具
将 .docx 文件转换为纯文本输出

用法: python read_word.py <input.docx>
输出: 文档内容打印到 stdout
"""

import sys
import os
import io

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx is not installed.", file=sys.stderr)
    print("Install it with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def read_docx(file_path: str) -> str:
    """读取 .docx 文件并返回文本内容"""
    try:
        doc = Document(file_path)

        # 提取段落文本
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 跳过空段落
                paragraphs.append(text)

        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        return '\n'.join(paragraphs)

    except Exception as e:
        print(f"ERROR: Failed to read document: {e}", file=sys.stderr)
        sys.exit(1)


def main():
    # 设置 stdout 为 UTF-8 编码，避免 Windows GBK 编码问题
    if sys.platform == 'win32':
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    if len(sys.argv) < 2:
        print("Usage: python read_word.py <input.docx>", file=sys.stderr)
        print("Output: Document content will be printed to stdout", file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]

    if not os.path.exists(input_path):
        print(f"ERROR: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    if not input_path.lower().endswith('.docx'):
        print(f"WARNING: Input file may not be a .docx file: {input_path}", file=sys.stderr)

    # 读取并输出文档内容
    content = read_docx(input_path)
    print(content)


if __name__ == "__main__":
    main()
